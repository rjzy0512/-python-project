# 使用import导入os模块
import os
# 使用import导入shutil模块
import shutil
# 使用import导入python-docx模块
import docx

# 将调查表的路径 /Users/yequ/调查表 赋值给变量preferPath
preferPath = "/Users/yequ/调查表"
# 用os.listdir()函数获取该路径下所有的调查表，并赋值给变量allItems
allItems = os.listdir(preferPath)

# 使用for循环遍历所有调查表
for item in allItems:
    # 使用os.path.join()函数拼接当前调查表的路径，并赋值给变量path
    path = os.path.join(preferPath, item)

    # TODO 使用docx.Document()函数读取当前调查表，并赋值给变量doc
    doc = docx.Document(path)

    # TODO 读取调查表中第三段的第二个样式块的文本内容【性别】，并赋值给变量gender
    gender = doc.paragraphs[2].runs[1].text

    # TODO 使用os.path.join()函数拼接新的文件夹路径，并赋值给变量genderPath
    genderPath = os.path.join(preferPath, gender)

    # TODO 使用if语句和os.path.exists()函数判断当该文件夹不存在时
    if not os.path.exists(genderPath):
        # TODO 如果不存在，使用os.mkdir()函数创建该文件夹
        os.mkdir(genderPath)

    # TODO 使用shutil.move()函数将调查表移动到对应的文件夹中
    shutil.move(path, genderPath)

    # TODO 使用os.path.join()函数拼接移动后的调查表的路径，并赋值给变量updatePath=
    updatePath = os.path.join(genderPath, item)

    # TODO 使用os.path.join()函数拼接新的文件夹"需要安静的学习环境"的路径，并赋值给变量quietPath
    quietPath = os.path.join(genderPath, "需要安静的学习环境")

    # TODO 使用if语句和os.path.exists()函数判断当该文件夹不存在时
    if not os.path.exists(quietPath):
        # TODO 如果不存在，使用os.mkdir()函数创建该文件夹
        os.mkdir(quietPath)

    # 使用os.path.join()函数拼接新的文件夹"不太需要安静的学习环境"的路径，并赋值给变量notQuietPath
    notQuietPath = os.path.join(genderPath, "不太需要安静的学习环境")
    # 使用if语句和os.path.exists()函数判断当该文件夹不存在时
    if not os.path.exists(notQuietPath):
        # 如果不存在，使用os.mkdir()函数创建该文件夹
        os.mkdir(notQuietPath)

    # TODO 使用 .tables 属性读取调查表中的第一个表格，并赋值给变量table
    table = doc.tables[0]

    # TODO 使用 .cell(row,col) 和 .text 属性读取表格中的第二项偏好分数，并赋值给变量score
    score = table.cell(2, 1).text

    # TODO 使用int()函数将变量score转换为整型，并赋值给变量scoreInt
    scoreInt = int(score)

    # TODO 使用if语句判断当第二项偏好的分数大于5时
    if scoreInt > 5:

        # TODO 如果分数大于5，使用shutil.move()函数将调查表移动到"需要安静的学习环境"的文件夹下
        shutil.move(updatePath, quietPath)

    # TODO 其他情况下，也就是当第二项偏好的分数小于等于5时
    else:

        # TODO 使用shutil.move()函数将文件移动到"不太需要安静的学习环境"的文件夹下
        shutil.move(updatePath, notQuietPath)
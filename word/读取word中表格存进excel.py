# 导入os模块
import os
# 导入docx
import docx
# 导入openpyxl 模块
import openpyxl

# 将报名表的路径 /Users/yequ/英语 赋值给变量allFormPath
allFormPath = r"C:\Users\Administrator\PycharmProjects\pythonProject\word\yequ\Downloads\registration\英语"
# 使用os.listdir()函数获取该路径下所有的班级文件夹，并赋值给变量allItems
allItems = os.listdir(allFormPath)

# TODO 定义一个空列表allStudentsData来装入所有学生的数据
allStudentsData = []

# TODO 使用for循环逐个遍历所有报名表

for baomingbiao in allItems:
    # TODO 定义一个空字典studentData装入单个学生数据
    studentData = {}

    # TODO 使用os.path.join()函数拼接出报名表路径，并赋值给变量keyPath=
    keyPath = os.path.join(allFormPath, baomingbiao)

    # TODO 使用docx.Document()函数读取报名表，并赋值给变量doc
    doc = docx.Document(keyPath)

    # TODO 读取报名表中的第一个表格，并赋值给变量table
    table = doc.tables[0]

    # 依次读取表格中的信息，并存入到studentData字典里的不同键中
    studentData["name"] = table.cell(0, 1).text
    studentData["gender"] = table.cell(0, 3).text
    studentData["hometown"] = table.cell(0, 5).text
    studentData["school"] = table.cell(1, 1).text
    studentData["grade"] = table.cell(1, 3).text
    # TODO 读取表格中志愿者的专业，并存入到studentData字典里的major键中
    studentData["major"] = table.cell(1, 5).text

    # TODO 读取表格中志愿者的联系电话，并存入到studentData字典里的tel键中
    studentData["tel"] = table.cell(2, 1).text

    studentData["qq"] = table.cell(2, 3).text
    studentData["yesorno"] = table.cell(2, 5).text

    # TODO 使用append()函数将studentData添加到allStudentsData列表中
    allStudentsData.append(studentData)

# 使用openpyxl.load_workbook()函数打开 /Users/yequ/志愿者统计.xlsx，并赋值给变量wb
wb = openpyxl.load_workbook(r"C:\Users\Administrator\PycharmProjects\pythonProject\word\yequ\Downloads\registration\志愿者统计.xlsx")
# 将名为 英语 的工作表赋值给sheet变量
sheet = wb["英语"]

# TODO 使用for循环和enumerate()函数
# 遍历所有学生数据allStudentsData的同时
# 生成一个从2开始的索引index
# 因为工作表中的第一行是表头并且在Excel中行数从1开始，所以我们的信息要从表格中的第二行开始填写

for index, studentData in enumerate(allStudentsData, 2):
    # 逐个将studentData字典中的信息填入到对应的单元格中
    sheet[f"A{index}"].value = studentData["name"]
    sheet[f"B{index}"].value = studentData["gender"]
    sheet[f"C{index}"].value = studentData["hometown"]
    sheet[f"D{index}"].value = studentData["school"]
    sheet[f"E{index}"].value = studentData["grade"]
    # TODO 将studentData字典中的学生专业填入到第F列中
    sheet[f"F{index}"].value = studentData["major"]

    # TODO 将studentData字典中的联系电话填入到第G列中
    sheet[f"G{index}"].value = studentData["tel"]

    sheet[f"H{index}"].value = studentData["qq"]
    sheet[f"I{index}"].value = studentData["yesorno"]

# 使用save()函数将工作簿保存到指定路径 /Users/yequ/志愿者统计.xlsx
wb.save(r"C:\Users\Administrator\PycharmProjects\pythonProject\word\yequ\Downloads\registration\志愿者统计.xlsx")
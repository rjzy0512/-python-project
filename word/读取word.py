# 使用import导入os模块
import os
# 使用import导入docx
import docx
# 使用import导入openpyxl模块
import openpyxl

# 第一大题选择题标准答案
standardOne = ['B', 'B', 'B', 'C', 'D', 'A', 'D', 'D', 'B', 'A']
# 第二大题填空题标准答案
standardTwo = ["东临碣石", "行舟绿水前", "孤山寺北贾亭西", "断肠人在天涯", "故人具鸡黍", "一曲新词酒一杯", "何当共剪西窗烛", "误入藕花深处", "烟笼寒水月笼沙", "万籁此都寂",
               "初日照高林", "腾蛇乘雾"]

# 将乔老师的答题卡文件夹路径 /Users/qiao/answerKey 赋值给变量allKeyPath
allKeyPath = r"C:\Users\Administrator\PycharmProjects\pythonProject\word\answerKey"
# 使用os.listdir()函数获取该路径下所有的文件，并赋值给变量allItems
allItems = os.listdir(allKeyPath)

# 定义一个空列表allStudentsData存储所有学生数据
allStudentsData = []

# 使用for循环逐个遍历所有学生答题卡
for item in allItems:
    # 定义一个空字典studentData存储单个学生数据
    studentData = {}

    # 使用os.path.splitext()函数获取文件名的前半段，并赋值给变量fileName
    fileName = os.path.splitext(item)[0]
    # 使用split()函数以"-"分隔文件名，将第1部分班级信息赋值到学生数据字典的classInfo键里
    studentData["classInfo"] = fileName.split("-")[0]
    # 使用split()函数以"-"分隔文件名，将第2部分姓名信息赋值到学生数据字典的name键里
    studentData["name"] = fileName.split("-")[1]

    # 使用os.path.join()函数拼接出答题卡路径，并赋值给变量keyPath
    keyPath = os.path.join(allKeyPath, item)
    # 读取答题卡并赋值给变量doc
    doc = docx.Document(keyPath)

    # 读取第四段学号段，并赋值给变量idPara
    idPara = doc.paragraphs[3]
    # 读取学号段中第二个样式块，并赋值给变量idRun
    idRun = idPara.runs[1]
    # 读取学号，并赋值到学生数据字典的id键里
    studentData["id"] = idRun.text

    # 初始化学生数据字典里scoreTwo字段为0分，作为填空题分数
    studentData["scoreTwo"] = 0

    # 使用for循环和enumerate()函数
    # 遍历储存标准答案的列表standardTwo的同时
    # 生成一个从8开始的idx
    for idx, value in enumerate(standardTwo, 8):
        # 获取学生答案，并赋值给变量studentAnswerTwo
        studentAnswerTwo = doc.paragraphs[idx].runs[1].text
        # 判断当学生答案与标准答案相等时
        if studentAnswerTwo == value:
            # 当学生答案与标准答案一样时，学生的填空题分数加5分
            studentData["scoreTwo"] = studentData["scoreTwo"] + 5

    # 读取答题卡中的表格，并赋值给变量table
    table = doc.tables[0]
    # 读取表格中的第二行，并赋值给变量secRow
    secRow = table.rows[1]
    # 读取第二行中的所有单元格，并赋值给变量secCell
    secCell = secRow.cells

    # TODO 初始化学生数据字典里scoreOne字段为0分，作为选择题分数
    studentData["scoreOne"] = 0

    # TODO 使用for循环和enumerate()函数
    # 遍历储存标准答案的列表standardOne的同时
    # 生成一个从0开始的idx
    for idx, value in enumerate(standardOne):

        # TODO 获取学生选择题答案，并赋值给变量studentAnswerOne
        studentAnswerOne = secCell[idx].text

        # TODO 判断当学生答案与标准答案相等时
        if studentAnswerOne == value:
            # TODO 当学生答案与标准答案一样时，学生的选择题分数加5分
            studentData["scoreOne"] += 5

    # 使用append()函数将studentData添加到总学生数据allStudentsData中
    allStudentsData.append(studentData)

# 创建一个新工作簿，并赋值给变量newWb
newWb = openpyxl.Workbook()
# 将名为Sheet的默认工作表赋值给变量sheet
sheet = newWb["Sheet"]
# 将sheet工作表名称修改为"成绩表"
sheet.title = "成绩表"

# 给工作表设置表头
sheet["A1"].value = "学生姓名"
sheet["B1"].value = "学生班级"
sheet["C1"].value = "学生学号"
sheet["D1"].value = "选择题分数"
sheet["E1"].value = "填空题分数"

# 使用for循环和enumerate()函数
# 遍历所有学生数据allStudentsData的同时
# 生成一个从2开始的index
for index, studentData in enumerate(allStudentsData, 2):
    # TODO 将第A列数据设置为学生姓名studentData["name"]
    sheet[f"A{index}"].value = studentData["name"]

    # 设置学生班级
    sheet[f"B{index}"].value = studentData["classInfo"]
    # 设置学生学号
    sheet[f"C{index}"].value = studentData["id"]
    # 设置选择题分数
    sheet[f"D{index}"].value = studentData["scoreOne"]
    # 设置填空题分数
    sheet[f"E{index}"].value = studentData["scoreTwo"]

# TODO 使用save()函数将工作簿保存到 /Users/qiao/answerKey/成绩汇总.xlsx 路径下
newWb.save(r"C:\Users\Administrator\PycharmProjects\pythonProject\word\answerKey\成绩汇总.xlsx")